# -*- coding: utf-8 -*-
"""
Created on Wed Mar  5 10:25:39 2025

@author: quang.truong
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
import open_directory

folder_path = open_directory.select_folder()

def report_output(result, folder_path):
    #prep dataframe
    result = result[result.Detail != ""]
    df_activity = result[result.Category == "Activity/Success"]
    df_pending = result[result.Category == "Pending Actions"]
    
    new_result = pd.DataFrame()
    for staff in df_activity["Staff"].unique():
        staff_result = df_activity[df_activity["Staff"] == staff]
        last_row = staff_result.iloc[-1].copy()
        last_row["Detail"] = ""  # Blank out the Detail column
        staff_result = pd.concat([staff_result, pd.DataFrame([last_row])], ignore_index=True)
        new_result = pd.concat([new_result, staff_result], ignore_index=True)
    df_activity = new_result
    
    new_result = pd.DataFrame()
    for staff in df_pending["Staff"].unique():
        staff_result = df_pending[df_pending["Staff"] == staff]
        last_row = staff_result.iloc[-1].copy()
        last_row["Detail"] = ""  # Blank out the Detail column
        staff_result = pd.concat([staff_result, pd.DataFrame([last_row])], ignore_index=True)
        new_result = pd.concat([new_result, staff_result], ignore_index=True)
    df_pending = new_result
    
    # List of tasks from the DataFrame
    tasks = df_activity["Detail"].tolist()
    pendings = df_pending.Detail.tolist()
    
    # Load the Template Word document
    #doc_path = r"C:\Users\quang.truong\OneDrive - HHS Office of the Secretary\Desktop\Report Generator\Template.docx"
    
    doc_path = folder_path + r"\Template.docx"
    doc = Document(doc_path)
    
    # Find the second column in the document (assuming a table exists)
    tables = doc.tables
    if tables:
        table = tables[0]  # Assuming we are modifying the first table
        cell = table.rows[1].cells[2]  # second row, middle column
        
        # Clear existing content
        cell.paragraphs[0].clear()
        
        # Add each task as a properly formatted bulleted list or a line break if empty
        for task in tasks:
            p = cell.add_paragraph()
            if task.strip():  # If task is not empty, add a bullet point
                p.style = "ListBullet"
                run = p.add_run(task)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            else:  # If empty, add just a blank line
                p.add_run("\n")
    
        # add pending items to consolidated report
        cell = table.rows[1].cells[4]  # second row, middle column
        
        # Clear existing content
        cell.paragraphs[0].clear()
        
        # Add each task as a properly formatted bulleted list or a line break if empty
        for pending in pendings:
            p = cell.add_paragraph()
            if pending.strip():  # If pending is not empty, add a bullet point
                p.style = "ListBullet"
                run = p.add_run(pending)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            else:  # If empty, add just a blank line
                p.add_run("\n")
    
    # Save the modified document
#    output_path = folder_path + r"\Compiled Report.docx"
#    doc.save(output_path)
    doc.save("Compiled Report.docx")    
# report_output(result, doc_path, output_path)
